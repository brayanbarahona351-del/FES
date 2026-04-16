import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import datetime

# --- CONFIGURACIÓN DE ALTO NIVEL ---
st.set_page_config(page_title="FES Moos Suite - Sanidad Policial", page_icon="🧠", layout="wide")

# --- ESTILO VISUAL DINÁMICO ---
st.markdown("""
    <style>
    .main { background-color: #f4f7f6; }
    .excel-header { 
        background: linear-gradient(135deg, #0B3B60 0%, #1E88E5 100%); 
        color: white; padding: 25px; text-align: center; border-radius: 10px; 
        box-shadow: 0 5px 15px rgba(0,0,0,0.2); margin-bottom: 25px;
    }
    .seccion-titulo { color: #0B3B60; border-bottom: 2px solid #1E88E5; padding-bottom: 5px; margin-top: 30px; margin-bottom: 20px;}
    .card-analisis { 
        background-color: white; padding: 30px; border-radius: 10px; 
        border-top: 5px solid #1E88E5; box-shadow: 0 4px 12px rgba(0,0,0,0.1); 
        margin-bottom: 25px; 
    }
    .card-recomendaciones { 
        background-color: #F8FDF9; padding: 30px; border-radius: 10px; 
        border-left: 5px solid #2E7D32; box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        margin-bottom: 25px; 
    }
    .stDataFrame { border-radius: 10px; overflow: hidden; box-shadow: 0 4px 8px rgba(0,0,0,0.05); }
    </style>
    """, unsafe_allow_html=True)

# --- BASE DE DATOS: 90 PREGUNTAS LITERALES ---
BANCO_FES = {
    1: ("En mi familia nos ayudamos y apoyamos realmente unos a otros", "CO", "V"),
    2: ("Los miembros de la familia guardan, a menudo, sentimientos para sí mismos", "EX", "F"),
    3: ("En nuestra familia discutimos mucho", "CT", "V"),
    4: ("En mi familia no hay muchas posibilidades de realizar actividades por propia iniciativa", "AU", "F"),
    5: ("En mi familia es muy importante tener éxito en lo que se hace", "AC", "V"),
    6: ("A menudo hablamos de temas políticos o sociales", "IC", "V"),
    7: ("Dedicamos mucho tiempo a las diversiones y al ocio", "SR", "V"),
    8: ("En mi familia no nos interesamos mucho por las actividades religiosas", "MR", "F"),
    9: ("En mi familia las tareas están claramente definidas y asignadas", "OR", "V"),
    10: ("En mi familia el cumplimiento de las reglas es muy estricto", "CN", "V"),
    11: ("A menudo nos vemos obligados a ocultar nuestros sentimientos", "EX", "F"),
    12: ("Nos sentimos libres de expresar nuestra ira", "CT", "V"),
    13: ("En nuestra casa no hay reglas rígidas", "CN", "F"),
    14: ("Damos mucha importancia a la limpieza y al orden", "OR", "V"),
    15: ("Cada uno de nosotros tiene sus propias metas y ambiciones", "AC", "V"),
    16: ("Nos gusta asistir a conciertos o exposiciones", "IC", "V"),
    17: ("Rara vez salimos a comer o de viaje juntos", "SR", "F"),
    18: ("En mi familia no se permite cuestionar la autoridad de los padres", "CN", "V"),
    19: ("Nos sentimos muy unidos como familia", "CO", "V"),
    20: ("En mi familia rara vez se discute", "CT", "F"),
    21: ("En mi familia se nos anima a ser independientes", "AU", "V"),
    22: ("Para nosotros, el éxito es más importante que la ayuda mutua", "AC", "V"),
    23: ("A menudo vamos a la biblioteca o leemos libros", "IC", "V"),
    24: ("Casi nunca tenemos invitados en casa", "SR", "F"),
    25: ("Rara vez rezamos o vamos a la iglesia juntos", "MR", "F"),
    26: ("En mi casa se cambia de opinión a menudo", "OR", "F"),
    27: ("En mi familia nos castigan si rompemos las reglas", "CN", "V"),
    28: ("Rara vez nos ayudamos unos a otros", "CO", "F"),
    29: ("Hablamos abiertamente de nuestros problemas", "EX", "V"),
    30: ("Casi nunca nos peleamos", "CT", "F"),
    31: ("En mi familia cada uno hace lo que quiere", "AU", "V"),
    32: ("Damos mucha importancia a ganar en los juegos o deportes", "AC", "V"),
    33: ("Nos interesan mucho las actividades culturales", "IC", "V"),
    34: ("A menudo vamos al cine o a eventos deportivos", "SR", "V"),
    35: ("Tenemos valores morales y religiosos muy claros", "MR", "V"),
    36: ("La puntualidad es muy importante en mi familia", "OR", "V"),
    37: ("En mi familia los padres son muy dominantes", "CN", "V"),
    38: ("Peleamos mucho por nimiedades", "CT", "V"),
    39: ("En mi familia la unión es lo primero", "CO", "V"),
    40: ("Ocultamos nuestras opiniones para no molestar", "EX", "F"),
    41: ("A menudo nos movemos por impulsos", "CN", "V"),
    42: ("Damos mucha importancia a la educación y a los estudios", "AC", "V"),
    43: ("Casi nunca discutimos temas intelectuales", "IC", "F"),
    44: ("Nuestros amigos son importantes para la familia", "SR", "V"),
    45: ("En mi familia no somos muy religiosos", "MR", "F"),
    46: ("A menudo llegamos tarde a nuestras citas familiares", "OR", "F"),
    47: ("Rara vez nos ayudamos unos a otros en las tareas", "CO", "F"),
    48: ("Decimos lo que pensamos sin miedo", "EX", "V"),
    49: ("En mi familia hay un ambiente muy pacífico", "CT", "F"),
    50: ("A menudo pedimos permiso antes de hacer cualquier cosa", "AU", "F"),
    51: ("Nos esforzamos mucho por ser los mejores", "AC", "V"),
    52: ("No nos gustan mucho las actividades artísticas", "IC", "F"),
    53: ("Dedicamos mucho tiempo a visitar amigos", "SR", "V"),
    54: ("Tenemos una moralidad muy estricta", "MR", "V"),
    55: ("En mi casa las cosas están siempre en su lugar", "OR", "V"),
    56: ("Es difícil que alguien pierda los estribos", "CN", "F"),
    57: ("Si alguien llega tarde, nos preocupamos por él", "CO", "V"),
    58: ("A veces es mejor no hablar para no pelear", "EX", "F"),
    59: ("En mi familia siempre hay alguien peleando", "CT", "V"),
    60: ("Rara vez tomamos decisiones por nuestra cuenta", "AU", "F"),
    61: ("Nos da igual ganar o perder en un juego", "AC", "F"),
    62: ("Nos gusta aprender cosas nuevas", "IC", "V"),
    63: ("Nos divertimos mucho cuando salimos juntos", "SR", "V"),
    64: ("Nuestras creencias religiosas son muy importantes", "MR", "V"),
    65: ("A menudo cambiamos nuestros planes familiares", "OR", "F"),
    66: ("En mi familia hay muchas órdenes", "CN", "V"),
    67: ("Nos apoyamos en los momentos difíciles", "CO", "V"),
    68: ("Rara vez compartimos nuestras preocupaciones", "EX", "F"),
    69: ("Casi siempre estamos discutiendo", "CT", "V"),
    70: ("Somos muy independientes los unos de los otros", "AU", "V"),
    71: ("Trabajamos duro para salir adelante", "AC", "V"),
    72: ("Casi nunca vamos al teatro", "IC", "F"),
    73: ("Nos gusta pasar tiempo con otras familias", "SR", "V"),
    74: ("Cumplimos las normas morales de la sociedad", "MR", "V"),
    75: ("Llevamos una vida muy organizada", "OR", "V"),
    76: ("Nadie se sale de lo que está mandado", "CN", "V"),
    77: ("Estamos siempre dispuestos a ayudarnos", "CO", "V"),
    78: ("No nos contamos nuestras cosas", "EX", "F"),
    79: ("En mi familia no hay peleas importantes", "CT", "F"),
    80: ("Rara vez hacemos algo por iniciativa propia", "AU", "F"),
    81: ("El éxito material es muy importante para nosotros", "AC", "V"),
    82: ("Nos gusta leer periódicos y revistas", "IC", "V"),
    83: ("Rara vez participamos en actividades sociales", "SR", "F"),
    84: ("No nos preocupa mucho la religión", "MR", "F"),
    85: ("Si se ensucia algo, se limpia en seguida", "OR", "V"),
    86: ("Cada uno hace lo que le da la gana", "CN", "F"),
    87: ("Somos una familia muy unida", "CO", "V"),
    88: ("Nos expresamos con mucha libertad", "EX", "V"),
    89: ("Las discusiones terminan en peleas", "CT", "V"),
    90: ("En mi casa las reglas son flexibles", "CN", "F")
}

# --- ESTRUCTURA SEGÚN MANUAL ---
JERARQUIA = {
    "RELACIONES": {
        "CO": ("Cohesión", "Grado en que los miembros se apoyan y ayudan entre sí."),
        "EX": ("Expresividad", "Libertad para actuar y expresar sentimientos abiertamente."),
        "CT": ("Conflicto", "Grado en que se expresan abiertamente la cólera y agresividad.")
    },
    "DESARROLLO": {
        "AU": ("Autonomía", "Independencia en la toma de decisiones."),
        "AC": ("Actuación", "Orientación de las actividades al éxito y competencia."),
        "IC": ("Intelectual-Cultural", "Interés por actividades políticas, sociales y culturales."),
        "SR": ("Social-Recreativo", "Participación en actividades de ocio y sociales."),
        "MR": ("Moralidad-Religiosidad", "Énfasis en valores éticos y religiosos.")
    },
    "ESTABILIDAD": {
        "OR": ("Organización", "Importancia del orden y planificación en el hogar."),
        "CN": ("Control", "Grado de sujeción a reglas y procedimientos fijos.")
    }
}

# --- FUNCIONES CLÍNICAS ---
def calcular_puntuaciones(respuestas):
    raw_scores = {sigla: 0 for subs in JERARQUIA.values() for sigla in subs.keys()}
    for i, (txt, sigla, clave) in BANCO_FES.items():
        if respuestas.get(i) == clave:
            raw_scores[sigla] += 1
            
    # Conversión Simulada a T-Scores
    t_scores = {s: min(max(int((p * 7.5) + 20), 0), 100) for s, p in raw_scores.items()}
    return raw_scores, t_scores

def nivel_pd(pd_val):
    if pd_val >= 8: return "Muy Alto"
    if pd_val >= 6: return "Alto"
    if pd_val >= 4: return "Medio"
    if pd_val >= 2: return "Bajo"
    return "Muy Bajo"

def analizar_tipologia_familiar(raw):
    if raw["AC"] >= 7 and raw["CO"] >= 5:
        return "Familia Orientada al Logro", "Se concluye que el sistema familiar es un entorno protector, definido como una 'Familia Orientada al Logro'. Existe alta motivación por el desarrollo personal y profesional, sostenida por una base afectiva sólida."
    elif raw["CT"] >= 7 and raw["CO"] <= 4:
        return "Familia con Dinámica de Conflicto", "Se concluye que el sistema familiar presenta una dinámica de alta tensión y conflicto. La falta de cohesión agrava la vulnerabilidad de sus miembros ante el estrés externo."
    elif raw["CN"] >= 7 and raw["OR"] >= 7:
        return "Familia Rígidamente Estructurada", "Se concluye que el hogar opera bajo el modelo de 'Familia Estructurada'. El orden y la disciplina son pilares fundamentales, aunque podrían estar limitando la expresividad espontánea."
    elif raw["CO"] >= 7 and raw["SR"] >= 6:
        return "Familia Integrada y Sociable", "Se concluye que el sistema es altamente funcional, definido como 'Familia Integrada'. Cuentan con excelentes redes de apoyo interno y externo."
    else:
        return "Familia en Desarrollo de Adaptación", "Se concluye que el sistema familiar se encuentra en una fase de adaptación, mostrando fortalezas y áreas de mejora mixtas dependientes del estresor actual o de la etapa del ciclo vital."

def generar_narrativa_dimensiones(raw):
    # A. RELACIONES
    rel_co = "alta cohesión emocional. Existe un fuerte sentimiento de pertenencia y apoyo mutuo entre los miembros." if raw["CO"] >= 6 else ("baja cohesión emocional, indicando un distanciamiento y desvinculación." if raw["CO"] <= 3 else "cohesión emocional promedio, con un apoyo mutuo funcional.")
    rel_ex = "La expresividad es alta, lo que indica que se permite la comunicación de sentimientos de manera abierta." if raw["EX"] >= 6 else ("La expresividad es limitada, sugiriendo dificultad para la comunicación abierta de sentimientos." if raw["EX"] <= 3 else "La expresividad es moderada, permitiendo comunicación en áreas seguras.")
    rel_ct = "El nivel de conflicto es elevado, sugiriendo un entorno de tensión y discusiones constantes." if raw["CT"] >= 6 else ("El nivel de conflicto es mínimo, sugiriendo un entorno pacífico y de resolución constructiva." if raw["CT"] <= 3 else "El conflicto se maneja dentro de los parámetros habituales.")
    texto_a = f"El/la evaluado/a percibe un clima familiar de {rel_co} {rel_ex} {rel_ct}"

    # B. DESARROLLO
    des_ac = f"altamente orientada a la actuación y el éxito. Existe una presión significativa por cumplir metas (Actuación: {raw['AC']}/9)." if raw["AC"] >= 6 else f"con expectativas de éxito equilibradas (Actuación: {raw['AC']}/9)."
    des_ocio = f"Esto influye en una disminución de actividades Sociales-Recreativas y Culturales, las cuales se perciben como secundarias o descuidadas frente a las obligaciones." if (raw["SR"] <= 4 and raw["AC"] >= 6) else "Mantienen un sano interés en actividades de ocio y cultura como complemento a sus obligaciones."
    texto_b = f"El perfil muestra una familia {des_ac} {des_ocio}"

    # C. ESTABILIDAD
    est_or = "sólida. Las tareas y responsabilidades están bien definidas y planificadas." if raw["OR"] >= 6 else "flexible o inestructurada, con baja planificación de la rutina diaria."
    est_cn = "El control es alto, denotando un sistema de reglas estricto que podría llegar a ser autoritario." if raw["CN"] >= 7 else ("El control es moderado, lo que indica que existen reglas claras pero estas no llegan a ser autoritarias ni rígidas, permitiendo un margen de libertad personal." if raw["CN"] >= 4 else "El control es bajo, sugiriendo alta permisividad en el hogar.")
    texto_c = f"El hogar presenta una organización {est_or} {est_cn}"

    # RECOMENDACIONES BASE
    recs = []
    if raw["AC"] >= 7: recs.extend(["Fomentar espacios de ocio y recreación familiar que no estén ligados a la productividad o el estudio.", "Equilibrar las demandas de 'Actuación' con intereses culturales o artísticos para un desarrollo integral."])
    if raw["CT"] >= 6: recs.append("Implementar técnicas de resolución pacífica de conflictos para evitar escaladas de tensión en el hogar.")
    if raw["CO"] <= 3: recs.append("Establecer rutinas de conexión familiar (Ej. cenas sin dispositivos, salidas conjuntas) para fortalecer el vínculo afectivo.")
    if raw["EX"] <= 3: recs.append("Promover la asertividad y la validación emocional, permitiendo que cada miembro exprese sus frustraciones sin temor a juicios.")
    if not recs: recs.append("Mantener las pautas de crianza y comunicación actuales, reforzando positivamente el apoyo mutuo.")
    
    return texto_a, texto_b, texto_c, recs

# --- INTERFAZ ---
if 'respuestas' not in st.session_state:
    st.session_state.respuestas = {i: None for i in range(1, 91)}

st.markdown('<div class="excel-header"><h1>ESCALA DE CLIMA SOCIAL FAMILIAR (FES)</h1><h3>Suite de Evaluación Pericial - Sanidad Policial</h3></div>', unsafe_allow_html=True)

with st.sidebar:
    st.header("👤 Datos Generales")
    nombre = st.text_input("Paciente / Evaluado", "Funcionario Policial X")
    edad = st.number_input("Edad", 18, 80, 30)
    ocup = st.text_input("Ocupación / Rango", "Sub-Inspector")
    lugar = st.text_input("Sede / Jurisdicción", "Sanidad Policial - Honduras")
    exam = st.text_input("Examinador", "Lic. en Psicología")
    fecha = st.date_input("Fecha de Evaluación", datetime.date.today())
    st.divider()

tab_test, tab_results = st.tabs(["📝 1. CUESTIONARIO FES", "📊 2. RESULTADOS E INFORME"])

with tab_test:
    st.markdown("<h2 class='seccion-titulo'>1. CUESTIONARIO (90 ÍTEMS)</h2>", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    for i, (txt, sub, clv) in BANCO_FES.items():
        col = c1 if i <= 45 else c2
        with col:
            st.session_state.respuestas[i] = st.radio(
                f"**{i}.** {txt}", 
                ["V", "F"], 
                key=f"q{i}", 
                horizontal=True, 
                index=None if st.session_state.respuestas[i] is None else ["V", "F"].index(st.session_state.respuestas[i])
            )
            st.divider()

if None not in st.session_state.respuestas.values():
    raw_scores, pt_scores = calcular_puntuaciones(st.session_state.respuestas)
    txt_a, txt_b, txt_c, recomendaciones = generar_narrativa_dimensiones(raw_scores)
    tipo_titulo, conclusion_narrativa = analizar_tipologia_familiar(raw_scores)

    with tab_results:
        # VISUALIZACIÓN STREAMLIT: HOJA DE RESPUESTAS MATRIZ
        st.markdown("<h2 class='seccion-titulo'>2. HOJA DE RESPUESTAS LLENA</h2>", unsafe_allow_html=True)
        resp_matrix = []
        for row in range(18): 
            fila = {}
            for col in range(5):
                item_num = row + 1 + (col * 18)
                fila[f"Ítem_{col+1}"] = item_num
                fila[f"R_{col+1}"] = st.session_state.respuestas[item_num]
            resp_matrix.append(fila)
        
        df_resp = pd.DataFrame(resp_matrix)
        st.dataframe(df_resp, hide_index=True, use_container_width=True)

        # VISUALIZACIÓN STREAMLIT: TABLA DE PUNTUACIONES
        st.markdown("<h2 class='seccion-titulo'>3. TABLA DE PUNTUACIONES Y PERFIL</h2>", unsafe_allow_html=True)
        data_puntos = []
        for dim, subs in JERARQUIA.items():
            for sigla, (nom, desc) in subs.items():
                pd_val = raw_scores[sigla]
                data_puntos.append({"Dimensión": dim, "Subescala": nom, "PD": pd_val, "Nivel": nivel_pd(pd_val)})
        
        df_puntos = pd.DataFrame(data_puntos)
        st.dataframe(df_puntos, hide_index=True, use_container_width=True)

        # ----------------------------------------------------------------------
        # SOLUCIÓN DEL ERROR TYPEERROR EN LA GRÁFICA
        # ----------------------------------------------------------------------
        names_full = []
        values_t = []
        colors_dim = []
        
        for dim, subs in JERARQUIA.items():
            for sigla, (nom, desc) in subs.items():
                names_full.append(nom)
                values_t.append(pt_scores[sigla])
                if dim == "RELACIONES": colors_dim.append("#1E88E5")
                elif dim == "DESARROLLO": colors_dim.append("#2E7D32")
                elif dim == "ESTABILIDAD": colors_dim.append("#E65100")

        fig = go.Figure(data=[go.Bar(x=names_full, y=values_t, marker_color=colors_dim)])
        fig.update_layout(yaxis_range=[0, 100], title="Perfil de T-Scores", height=350, margin=dict(t=40, b=0))
        st.plotly_chart(fig, use_container_width=True)
        # ----------------------------------------------------------------------

        # VISUALIZACIÓN STREAMLIT: INTERPRETACIÓN DE RESULTADOS IA
        st.markdown(f"""
        <div class="card-analisis">
            <h2>4. INTERPRETACIÓN DE RESULTADOS</h2>
            <h4>A. Dimensión de Relaciones</h4>
            <p>{txt_a}</p>
            <h4>B. Dimensión de Desarrollo</h4>
            <p>{txt_b}</p>
            <h4>C. Dimensión de Estabilidad</h4>
            <p>{txt_c}</p>
        </div>
        <div class="card-recomendaciones">
            <h2>5. CONCLUSIONES Y RECOMENDACIONES</h2>
            <p><b>Diagnóstico: {tipo_titulo}</b><br>{conclusion_narrativa}</p>
            <p><b>Recomendaciones:</b></p>
            <ul>{''.join([f'<li>{r}</li>' for r in recomendaciones])}</ul>
        </div>
        """, unsafe_allow_html=True)

        # --- GENERADOR DE WORD ---
        st.markdown("<h2 class='seccion-titulo'>💾 EXPORTACIÓN OFICIAL</h2>", unsafe_allow_html=True)
        
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # TÍTULO DEL DOCUMENTO
        doc.add_heading('INFORME CLÍNICO - ESCALA FES DE MOOS', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. FICHA TÉCNICA
        doc.add_heading('1. FICHA TÉCNICA', level=1)
        doc.add_paragraph(f"Paciente / Evaluado: {nombre}\nEdad: {edad} años\nOcupación: {ocup}\nSede / Jurisdicción: {lugar}\nExaminador: {exam}\nFecha de Evaluación: {fecha.strftime('%d/%m/%Y')}")

        # 2. HOJA DE RESPUESTAS
        doc.add_heading('2. HOJA DE RESPUESTAS LLENA', level=1)
        doc.add_paragraph("Instrucciones: El evaluado marcó V para Verdadero y F para Falso.")
        
        table_resp = doc.add_table(rows=19, cols=10)
        table_resp.style = 'Table Grid'
        
        for col_idx in range(5):
            table_resp.cell(0, col_idx*2).text = "Ítem"
            table_resp.cell(0, col_idx*2).paragraphs[0].runs[0].bold = True
            table_resp.cell(0, col_idx*2 + 1).text = "R"
            table_resp.cell(0, col_idx*2 + 1).paragraphs[0].runs[0].bold = True

        for row in range(18):
            for col in range(5):
                item_num = row + 1 + (col * 18)
                table_resp.cell(row+1, col*2).text = str(item_num)
                table_resp.cell(row+1, col*2 + 1).text = st.session_state.respuestas[item_num]

        doc.add_paragraph() 

        # 3. TABLA DE PUNTUACIONES Y PERFIL
        doc.add_heading('3. TABLA DE PUNTUACIONES Y PERFIL', level=1)
        doc.add_paragraph("(Conversión de Puntajes Directos a Niveles Interpretativos)")
        
        table_pts = doc.add_table(rows=1, cols=4)
        table_pts.style = 'Table Grid'
        hdr_cells = table_pts.rows[0].cells
        hdr_cells[0].text = 'Dimensión'
        hdr_cells[1].text = 'Subescala'
        hdr_cells[2].text = 'PD'
        hdr_cells[3].text = 'Nivel'
        for cell in hdr_cells: cell.paragraphs[0].runs[0].bold = True

        for r_data in data_puntos:
            row_cells = table_pts.add_row().cells
            row_cells[0].text = r_data["Dimensión"]
            row_cells[1].text = r_data["Subescala"]
            row_cells[2].text = str(r_data["PD"])
            row_cells[3].text = r_data["Nivel"]

        # 3.1 GRÁFICA EN WORD
        plt.figure(figsize=(9, 4))
        plt.bar(names_full, values_t, color=colors_dim)
        plt.axhline(y=50, color='red', linestyle='--', alpha=0.5)
        plt.ylim(0, 100)
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        img_buf = BytesIO()
        plt.savefig(img_buf, format='png')
        img_buf.seek(0)
        doc.add_picture(img_buf, width=Inches(6.0))
        plt.close()

        doc.add_page_break()

        # 4. INTERPRETACIÓN DE RESULTADOS
        doc.add_heading('4. INTERPRETACIÓN DE RESULTADOS', level=1)
        doc.add_heading('A. Dimensión de Relaciones', level=2)
        doc.add_paragraph(txt_a)
        doc.add_heading('B. Dimensión de Desarrollo', level=2)
        doc.add_paragraph(txt_b)
        doc.add_heading('C. Dimensión de Estabilidad', level=2)
        doc.add_paragraph(txt_c)

        # 5. CONCLUSIONES Y RECOMENDACIONES
        doc.add_heading('5. CONCLUSIONES Y RECOMENDACIONES', level=1)
        p_diag = doc.add_paragraph()
        p_diag.add_run(conclusion_narrativa)
        
        doc.add_paragraph("\nRecomendaciones Terapéuticas:")
        for r in recomendaciones:
            doc.add_paragraph(r, style='List Bullet')

        # Firmas
        doc.add_paragraph("\n\n\n" + "_"*40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_firma = doc.add_paragraph(f"{exam}\n{lugar}")
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # BOTÓN DE DESCARGA
        buf = BytesIO()
        doc.save(buf)
        st.download_button(
            label="📥 DESCARGAR INFORME CLÍNICO FINAL (WORD)", 
            data=buf.getvalue(), 
            file_name=f"Informe_FES_{nombre.replace(' ', '_')}.docx", 
            type="primary", 
            use_container_width=True
        )

else:
    with tab_results:
        st.info("⚠️ Complete las 90 preguntas en la pestaña del cuestionario para generar las gráficas, el análisis y el documento Word.")
